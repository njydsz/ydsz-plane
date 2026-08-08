#!/usr/bin/env python3
"""
用 python-docx 将精简补充内容合并回原 PRD（方式：克隆原文档 + 追充分节）
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell):
    """设置单元格边框"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_run_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    """设置 Run 的字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)


def add_heading_para(doc, text, level, color_rgb=None):
    """添加标题段落"""
    size_map = {0: 18, 1: 16, 2: 14, 3: 12}
    size = size_map.get(level, 14)
    color_rgb = color_rgb or (31, 78, 121)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color_rgb)
    # 段前段后
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_para(doc, text, size=10.5, bold=False, color=None, align=None, italic=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_code_para(doc, text):
    """添加代码/模型段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 灰色背景
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, headers, rows, header_color='1F4E79'):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 不设置 table.style，依赖手动设置的边框
    
    # 表头
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_shading(cell, header_color)
        set_cell_borders(cell)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            
            # 颜色标记优先级
            color_map = {'P0': (204, 0, 0), 'P1': (230, 126, 34), 'P2': (39, 174, 96)}
            if str(cell_text) in color_map:
                set_run_font(run, size=10, bold=True, color=color_map[str(cell_text)])
            else:
                set_run_font(run, size=10)
            
            set_cell_borders(cell)
            # 交替行
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'F5F7FA')
    
    return table


def add_horizontal_line(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(20)
    return p


def build_supplement(doc):
    """构建补充内容追加到 doc"""
    
    # ============ 页标题 ============
    doc.add_page_break()
    add_heading_para(doc, '补充章节：共性能力对标后新增模块', level=0, color_rgb=(31, 78, 121))
    
    add_para(doc,
        '本章基于国内头部厂商（云效、TAPD、ONES、PingCode）共性功能清单，对 Ydsz Plane PRD V1.0 进行对标后，'
        '聚焦于符合产品定位（轻量级开源 PM、中小敏捷团队）的核心缺失能力。'
        '排除 DevOps 交付链路、重量级测试管理等属于生态对接或商业版的能力。',
        color=(102, 102, 102), italic=True)
    
    # ============ 优先级总览 ============
    add_heading_para(doc, '新增模块优先级总览', level=1, color_rgb=(46, 117, 182))
    
    add_table(doc,
        ['模块', '功能域', '优先级', '原 PRD 状态'],
        [
            ['模块 N', '自定义字段体系', 'P0', '完全缺失'],
            ['模块 O', '可视化工作流引擎', 'P1', '仅部分（状态组）'],
            ['模块 P', '操作审计日志', 'P1', '仅部分（活动日志）'],
            ['模块 Q', '项目集管理', 'P1', '完全缺失'],
            ['模块 R', '通用审批流配置', 'P2', '未覆盖'],
        ])
    
    # ============ 模块 N ============
    doc.add_page_break()
    add_heading_para(doc, '模块 N：自定义字段体系（P0）', level=0)
    add_para(doc,
        '当前 PRD 属性仅描述固定字段，缺少自定义字段扩展能力。'
        '这是所有成熟项目管理工具的标配能力，对标 Jira Custom Fields 和云效自定义字段。',
        color=(102, 102, 102), italic=True)
    add_para(doc, '【适用范围】', bold=True)
    add_para(doc, '所有工作项类型（需求/任务/缺陷）和自定义工件，字段类型包括：文本、数字、单选、多选枚举、日期、人员、URL、公式、级联。')
    
    add_heading_para(doc, 'N.1 功能需求详述', level=2)
    add_table(doc,
        ['功能点', '功能描述', '优先级'],
        [
            ['字段类型', '文本、数字、单选枚举、多选枚举、日期、人员、URL、公式、级联', 'P0'],
            ['适用范围', '自定义字段可作用于需求/任务/缺陷等不同工件类型', 'P0'],
            ['字段配置', '字段名称、标识、占位提示、默认值、是否必填、是否唯一', 'P0'],
            ['枚举管理', '枚举字段支持预设选项列表，可排序、设颜色、启用/禁用', 'P0'],
            ['字段布局', '创建/编辑/详情页字段可拖拽排序，支持字段分组', 'P1'],
            ['字段权限', '按角色配置字段的可见性和编辑权限', 'P1'],
            ['全局/项目字段', '支持空间级全局字段和项目级独立字段', 'P1'],
            ['筛选分组排序', '自定义字段可参与过滤、分组、排序和报表统计', 'P1'],
        ])
    
    add_heading_para(doc, 'N.2 数据模型设计', level=2)
    add_code_para(doc, 'CustomField: workspace(FK), name, key, field_type, applicable_types(JSON),')
    add_code_para(doc, '    is_required, default_value, scope(global/project), created_at')
    add_code_para(doc, 'CustomFieldOption: field(FK), label, value, color, sort_order, is_active')
    add_code_para(doc, 'CustomFieldValue: field(FK), issue(FK), value_text, value_number, value_date, value_user')
    add_code_para(doc, 'FieldLayout: project(FK), applicable_type, fields(JSON ordered_list)')
    
    # ============ 模块 O ============
    doc.add_page_break()
    add_heading_para(doc, '模块 O：可视化工作流引擎（P1）', level=0)
    add_para(doc,
        '当前 PRD 状态组设计中仅有状态组概念，缺少可视化流程配置界面。'
        '对标 Jira Workflow Designer、ONES 状态机、云效流程引擎。',
        color=(102, 102, 102), italic=True)
    
    add_heading_para(doc, 'O.1 功能需求详述', level=2)
    add_table(doc,
        ['功能点', '功能描述', '优先级'],
        [
            ['可视化流程设计', '拖拽式画布配置状态节点和转换箭头，支持条件分支', 'P1'],
            ['状态转换规则', '配置状态间是否可流转，是否需审批/校验/填写必填字段', 'P1'],
            ['按类型独立配置', '需求/任务/缺陷可配置不同的工作流', 'P1'],
            ['触发器（Trigger）', '状态转换时触发自动动作（通知/字段变更/子工作项联动）', 'P1'],
            ['校验条件（Validator）', '流转前校验（如关闭任务时所有子任务必须已完成）', 'P2'],
            ['后置函数（Post Function）', '流转后自动执行动作（如流转开发中自动分配提交人）', 'P2'],
            ['版本管理', '工作流修改产生新版本，旧版本保留以兼容历史数据', 'P2'],
        ])
    
    add_heading_para(doc, 'O.2 数据模型设计', level=2)
    add_code_para(doc, 'Workflow: project(FK), name, applicable_type, version(INT), is_active, diagram_data(JSON)')
    add_code_para(doc, 'WorkflowState: workflow(FK), state(FK State), position_x, position_y')
    add_code_para(doc, 'WorkflowTransition: workflow(FK), name, from_state(FK), to_state(FK),')
    add_code_para(doc, '    conditions(JSON), validators(JSON), post_functions(JSON), require_approval(BOOLEAN)')
    
    # ============ 模块 P ============
    doc.add_page_break()
    add_heading_para(doc, '模块 P：操作审计日志（P1）', level=0)
    add_para(doc,
        '当前 PRD 仅提及模块内的活动日志，缺少全局独立审计模块。'
        '对标等保三级合规要求。',
        color=(102, 102, 102), italic=True)
    
    add_heading_para(doc, 'P.1 功能需求详述', level=2)
    add_table(doc,
        ['功能点', '功能描述', '优先级'],
        [
            ['全量操作留痕', '记录登录/设置变更/权限变更/数据修改/删除/导出等所有操作', 'P1'],
            ['审计维度', '操作人员、操作时间、类型、对象、IP、User-Agent、变更前后值', 'P1'],
            ['审计查询', '按人员、时间范围、操作类型、对象多维度筛选', 'P1'],
            ['审计导出', '支持审计日志导出 CSV/Excel 用于合规审查', 'P1'],
            ['敏感操作报警', '批量删除、权限提升、导出敏感数据触发即时告警', 'P2'],
            ['保留策略', '可配置日志保留时长，默认永久，合规场景设上限', 'P2'],
        ])
    
    add_heading_para(doc, 'P.2 数据模型设计', level=2)
    add_code_para(doc, 'AuditLog: workspace(FK), project(FK), user(FK), action_type, target_type, target_id,')
    add_code_para(doc, '    before_value(JSON), after_value(JSON), ip_address, user_agent, created_at')
    add_code_para(doc, 'AuditLogConfig: workspace(FK), retention_days(INT), alert_rules(JSON)')
    
    # ============ 模块 Q ============
    doc.add_page_break()
    add_heading_para(doc, '模块 Q：项目集管理（P1）', level=0)
    add_para(doc,
        '当前 PRD 仅支持单项目操作，缺少多项目聚合管控能力。'
        '对标 ONES 项目集、云效项目集、TAPD 项目组合。',
        color=(102, 102, 102), italic=True)
    
    add_heading_para(doc, 'Q.1 功能需求详述', level=2)
    add_table(doc,
        ['功能点', '功能描述', '优先级'],
        [
            ['项目集创建', '聚合多个关联项目，设置目标、负责人、关联业务线', 'P1'],
            ['项目集看板', '跨项目聚合进度（需求/任务/缺陷总数与完成率）、风险项一览', 'P1'],
            ['跨项目甘特图', '时间轴展示多个项目关键里程碑，支持依赖关系标记', 'P1'],
            ['资源统筹', '查看多项目下的人员负载分布，识别资源瓶颈', 'P2'],
            ['跨项目依赖', '标记项目 A 依赖项目 B 的版本发布，完成自动通知', 'P2'],
            ['项目集报告', '自动生成交付报告（多项目进度、质量横向对比）', 'P2'],
            ['项目集权限', '独立于单项目的权限，支持 PMO/VP 级别只读查看', 'P2'],
        ])
    
    add_heading_para(doc, 'Q.2 数据模型设计', level=2)
    add_code_para(doc, 'ProjectProgram: workspace(FK), name, description, owner(FK), status, created_at')
    add_code_para(doc, 'ProjectProgramMember: program(FK), user(FK), role(viewer/manager/editor)')
    add_code_para(doc, 'ProjectProgramRelation: program(FK), project(FK), added_at, added_by')
    add_code_para(doc, 'ProjectProgramDependency: program(FK), source_project(FK), target_project(FK),')
    add_code_para(doc, '    dependency_type, due_date, status')
    
    # ============ 模块 R ============
    doc.add_page_break()
    add_heading_para(doc, '模块 R：通用审批流配置（P2）', level=0)
    add_para(doc,
        '当前 PRD 仅在"文档评审"中提及审批节点，缺少通用的审批引擎，使所有工件类型都能灵活配置审批流程。',
        color=(102, 102, 102), italic=True)
    add_para(doc, '【核心场景】', bold=True)
    add_para(doc, '需求上线审批、缺陷关闭审批、自定义状态流转审批、跨项目交付物验收审批。')
    
    add_heading_para(doc, 'R.1 功能需求详述', level=2)
    add_table(doc,
        ['功能点', '功能描述', '优先级'],
        [
            ['多节点审批链', '可配置 1~N 级审批节点，每级指定审批人（指定人/角色/部门主管）', 'P2'],
            ['审批类型', '会签（全部通过）或签（任一通过）依次审批（逐人审批）', 'P2'],
            ['抄送规则', '审批通过/拒绝后抄送指定人员（如 PM、技术负责人）', 'P2'],
            ['触发条件', '按工件类型/状态变更/优先级等条件触发审批流', 'P2'],
            ['审批代理', '审批人可设置代理人（请假/出差时自动转交）', 'P2'],
            ['审批历史', '完整记录审批过程（审批人、意见、时间）', 'P2'],
            ['移动端审批', '审批人可在 IM/移动端通过/驳回（依赖 API 支持）', 'P2'],
        ])
    
    add_heading_para(doc, 'R.2 数据模型设计', level=2)
    add_code_para(doc, 'ApprovalFlow: project(FK), name, applicable_type, trigger_condition(JSON), is_active')
    add_code_para(doc, 'ApprovalNode: flow(FK), order(INT), approver_type(user/role/manager),')
    add_code_para(doc, '    approver_id, approval_type(and/or/sequential), cc_users(JSON)')
    add_code_para(doc, 'ApprovalRecord: issue(FK), flow(FK), node(FK), actor(FK),')
    add_code_para(doc, '    action(approve/reject), comment, created_at')
    add_code_para(doc, 'ApprovalDelegate: user(FK), delegate_user(FK), start_date, end_date, is_active')
    
    # ============ 实施路线图 ============
    doc.add_page_break()
    add_heading_para(doc, '附录：补充模块实施路线图', level=0)
    
    add_heading_para(doc, 'Phase 1（v1.1-v1.2）：基础扩展能力', level=2)
    add_para(doc, '• 模块 N 自定义字段体系 — 向后兼容，扩展 Issue 模型')
    add_para(doc, '• 模块 P 操作审计日志 — 独立模块，监听领域事件')
    
    add_heading_para(doc, 'Phase 2（v1.3-v1.5）：流程管控能力', level=2)
    add_para(doc, '• 模块 O 可视化工作流引擎 — 差异化竞争力，提升产品成熟度')
    add_para(doc, '• 模块 R 通用审批流配置 — 扩展 WorkflowTransition 审批能力')
    
    add_heading_para(doc, 'Phase 3（v1.6-v1.8）：组织扩展能力', level=2)
    add_para(doc, '• 模块 Q 项目集管理 — 面向 PMO/多团队场景，需完整 UI 支撑')
    
    # ============ 与现有模块关联 ============
    add_heading_para(doc, '与现有模块的关联关系', level=2)
    add_para(doc, '• 模块 N（自定义字段）影响所有工作项属性，需设计通用 EAV/JSON Schema 存储，向下兼容 Issue 表')
    add_para(doc, '• 模块 O（可视化工作流）扩展现有 State 模型，增加 Workflow 抽象层，兼容现有 6 状态组')
    add_para(doc, '• 模块 P（审计）独立模块，监听所有领域事件，不影响现有业务表结构')
    add_para(doc, '• 模块 Q（项目集）依赖现有 Project/Workspace 层级，Program 作为 Project 聚合层')
    add_para(doc, '• 模块 R（审批流）基于现有 WorkflowTransition，在 require_approval 标记时触发审批节点表')
    
    # ============ 结尾线 ============
    add_horizontal_line(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— Ydsz Plane PRD V1.0 补充章节（精简版）完 —')
    set_run_font(run, size=10, color=(153, 153, 153))
    run.italic = True


def main():
    original_path = r'D:\Code\open\ydsz-plane\docs\Ydsz Plane 产品需求文档.docx'
    backup_path = r'D:\Code\open\ydsz-plane\docs\Ydsz Plane 产品需求文档_v1.0_备份.docx'
    output_path = original_path  # 覆盖原文件
    
    # 用备份文件作为基础（因为当前原文件已经被 ElementTree 破坏了）
    print(f'使用备份文件: {backup_path}')
    
    # 打开原 docx 文档
    doc = Document(backup_path)
    print(f'已加载文档，段落数: {len(doc.paragraphs)}')
    
    # 追加补充内容
    build_supplement(doc)
    print(f'追加后段落数: {len(doc.paragraphs)}')
    
    # 保存
    doc.save(output_path)
    print(f'已保存到: {output_path}')
    
    # 输出文件大小信息
    original_size = os.path.getsize(backup_path)
    output_size = os.path.getsize(output_path)
    print(f'原文件大小: {original_size:,} bytes')
    print(f'输出文件大小: {output_size:,} bytes')
    print(f'增量: {output_size - original_size:,} bytes ({((output_size/original_size - 1) * 100):.1f}%)')
    print('\n✅ 完成！')


if __name__ == '__main__':
    main()
